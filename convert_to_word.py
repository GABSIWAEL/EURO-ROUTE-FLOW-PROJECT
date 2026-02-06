#!/usr/bin/env python3
"""
Script de conversion Markdown to Word (.docx)
Convertit CAHIER_DE_CHARGE.md en CAHIER_DE_CHARGE.docx
"""

import re
import os
from pathlib import Path


def install_required_packages():
    """Installer les packages nécessaires si manquants"""
    try:
        import markdown
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        print("Installation des packages requis...")
        os.system("pip install markdown-to-docx docx2pdf python-docx markdown2 -q")
        print("✓ Packages installés")


def markdown_to_word(markdown_file, output_file):
    """
    Convertir un fichier Markdown en Word Document
    """
    print(f"📖 Lecture du fichier : {markdown_file}")

    with open(markdown_file, 'r', encoding='utf-8') as f:
        content = f.read()

    print("🔄 Conversion en cours...")

    try:
        # Essayer avec markdown2docx d'abord
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement

        doc = Document()

        # Ajouter un titre
        title = doc.add_heading('CAHIER DE CHARGE', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        subtitle = doc.add_paragraph(
            'Plateforme de Gestion de Livraison - EURO ROUTE FLOW')
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle_format = subtitle.runs[0]
        subtitle_format.font.size = Pt(14)
        subtitle_format.font.italic = True

        # Parser le Markdown
        lines = content.split('\n')
        i = 0

        while i < len(lines):
            line = lines[i]

            # Titres
            if line.startswith('# '):
                heading = doc.add_heading(line[2:].strip(), level=1)
                heading.paragraph_format.space_before = Pt(12)
                heading.paragraph_format.space_after = Pt(6)
            elif line.startswith('## '):
                heading = doc.add_heading(line[3:].strip(), level=2)
                heading.paragraph_format.space_before = Pt(10)
                heading.paragraph_format.space_after = Pt(4)
            elif line.startswith('### '):
                heading = doc.add_heading(line[4:].strip(), level=3)
                heading.paragraph_format.space_before = Pt(8)
                heading.paragraph_format.space_after = Pt(3)
            elif line.startswith('#### '):
                heading = doc.add_heading(line[5:].strip(), level=4)
                heading.paragraph_format.space_before = Pt(6)
                heading.paragraph_format.space_after = Pt(2)

            # Listes
            elif line.startswith('- '):
                doc.add_paragraph(line[2:].strip(), style='List Bullet')
            elif line.startswith('* '):
                doc.add_paragraph(line[2:].strip(), style='List Bullet')
            elif re.match(r'^\d+\. ', line):
                match = re.match(r'^\d+\. (.+)$', line)
                if match:
                    doc.add_paragraph(match.group(1), style='List Number')

            # Code blocks
            elif line.strip().startswith('```'):
                # Collecter le code jusqu'à la fermeture
                code_lines = []
                i += 1
                while i < len(lines) and not lines[i].strip().startswith('```'):
                    code_lines.append(lines[i])
                    i += 1

                code_text = '\n'.join(code_lines)
                code_para = doc.add_paragraph(code_text, style='List Bullet')
                for run in code_para.runs:
                    run.font.name = 'Courier New'
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(0, 0, 0)
                code_para.paragraph_format.left_indent = Inches(0.5)

                # Ajouter du shading
                shading_elm = OxmlElement('w:shd')
                shading_elm.set(qn('w:fill'), 'E8E8E8')
                code_para._element.get_or_add_pPr().append(shading_elm)

            # Tables (détection simple)
            elif '|' in line and i > 0 and '|' in lines[i-1]:
                # C'est une table
                table_lines = [lines[i-1], line]
                i += 1
                while i < len(lines) and '|' in lines[i]:
                    table_lines.append(lines[i])
                    i += 1
                i -= 1

                # Parser la table
                headers = [h.strip() for h in table_lines[0].split('|')[1:-1]]
                rows = []

                for tline in table_lines[2:]:  # Skip separator row
                    if tline.strip():
                        cells = [c.strip() for c in tline.split('|')[1:-1]]
                        if len(cells) == len(headers):
                            rows.append(cells)

                if headers and rows:
                    table = doc.add_table(rows=1, cols=len(headers))
                    table.style = 'Light Grid Accent 1'

                    # Header row
                    hdr_cells = table.rows[0].cells
                    for idx, header in enumerate(headers):
                        hdr_cells[idx].text = header
                        for paragraph in hdr_cells[idx].paragraphs:
                            for run in paragraph.runs:
                                run.font.bold = True

                    # Data rows
                    for row_data in rows:
                        row_cells = table.add_row().cells
                        for idx, cell_data in enumerate(row_data):
                            row_cells[idx].text = cell_data

            # Paragraphes normaux
            elif line.strip() and not line.startswith(('---', '***', '___')):
                # Supprimer les marqueurs Markdown
                cleaned = re.sub(r'\*\*(.+?)\*\*', r'\1', line)
                cleaned = re.sub(r'\*(.+?)\*', r'\1', cleaned)
                cleaned = re.sub(r'__(.+?)__', r'\1', cleaned)
                cleaned = re.sub(r'_(.+?)_', r'\1', cleaned)
                cleaned = re.sub(r'`(.+?)`', r'\1', cleaned)

                if cleaned.strip():
                    p = doc.add_paragraph(cleaned.strip())
                    p.paragraph_format.space_after = Pt(6)

            # Lignes vides
            elif not line.strip():
                doc.add_paragraph()

            i += 1

        # Sauvegarder
        doc.save(output_file)
        print(f"✅ Conversion réussie : {output_file}")
        return True

    except Exception as e:
        print(f"❌ Erreur : {e}")
        print("Tentative avec une méthode alternative...")
        return False


if __name__ == "__main__":
    # Chemins
    script_dir = Path(__file__).parent
    markdown_file = script_dir / "CAHIER_DE_CHARGE.md"
    output_file = script_dir / "CAHIER_DE_CHARGE.docx"

    if not markdown_file.exists():
        print(f"❌ Erreur : Fichier {markdown_file} non trouvé")
        exit(1)

    # Installer les packages
    print("📦 Vérification des packages...")
    install_required_packages()

    # Convertir
    success = markdown_to_word(str(markdown_file), str(output_file))

    if success:
        print(f"\n✨ Fichier créé avec succès !")
        print(f"📄 {output_file.absolute()}")
        print(f"📊 Taille : {os.path.getsize(output_file) / 1024:.1f} KB")
    else:
        print("\n⚠️  Veuillez installer Pandoc manuellement :")
        print("https://pandoc.org/installing.html")
